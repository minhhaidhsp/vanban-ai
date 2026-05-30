"""
DOCX generation service — python-docx + lxml HTML parser.

Converts an Nd30Data dict into an A4 Word document following NĐ30/2020 formatting.
Fonts: Times New Roman 14pt for canCu/noiDung content, 12-13pt for header fields.
Margins: top/bottom 25mm, left 30mm, right 20mm  (NĐ30 standard).

HTML from TipTap is parsed with lxml to preserve:
  - Block structure: <p>, <h1-h6>, <ul>, <ol>
  - Inline formatting: <strong>, <em>, <u>, <s>, <span style="font-size/font-family">
  - Text alignment from inline style (text-align: center/right/justify/left)
"""
import asyncio
import io
import logging
import re
from typing import Any

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import html as lxml_html

logger = logging.getLogger(__name__)


# ── Inline run builder ────────────────────────────────────────────────────────

def _parse_inline(
    element, para,
    bold: bool = False, italic: bool = False,
    underline: bool = False, strike: bool = False,
    font_size: float = 14.0,
    font_family: str = "Times New Roman",
):
    """Recursively add text runs from *element*'s subtree to *para*.

    Inherits bold/italic/underline/strike/size/family from ancestor tags.
    """
    tag = getattr(element, "tag", None)
    b, i, u, s = bold, italic, underline, strike
    fs, ff = font_size, font_family

    if tag in ("strong", "b"):
        b = True
    elif tag in ("em", "i"):
        i = True
    elif tag == "u":
        u = True
    elif tag in ("s", "del", "strike"):
        s = True
    elif tag == "span":
        sty = element.get("style", "")
        m = re.search(r"font-size:\s*([\d.]+)pt", sty)
        if m:
            fs = float(m.group(1))
        m = re.search(r"font-family:\s*([^;]+)", sty)
        if m:
            first = m.group(1).strip().split(",")[0].strip().strip("'\"")
            if first:
                ff = first
    # <mark>, <code>, <a> etc. — pass through without special formatting

    def _add(text: str):
        if not text:
            return
        run = para.add_run(text)
        run.font.name = ff
        run.font.size = Pt(fs)
        run.font.bold = b
        run.font.italic = i
        if u:
            run.font.underline = True
        if s:
            run.font.strike = True

    _add(element.text)
    for child in element:
        _parse_inline(child, para, b, i, u, s, fs, ff)
        _add(child.tail)


# ── Block-level helpers ───────────────────────────────────────────────────────

_ALIGN_MAP: dict[str, WD_ALIGN_PARAGRAPH] = {
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
}


def _block_align(element, default: WD_ALIGN_PARAGRAPH) -> WD_ALIGN_PARAGRAPH:
    m = re.search(r"text-align:\s*(\w+)", element.get("style", ""))
    return _ALIGN_MAP.get(m.group(1), default) if m else default


def _new_para(doc: Document, align: WD_ALIGN_PARAGRAPH,
              space_before: float = 0.0, space_after: float = 0.0):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def _process_block(
    doc: Document, element,
    default_align: WD_ALIGN_PARAGRAPH,
    default_size: float,
    default_italic: bool,
):
    """Convert one block-level HTML element to docx paragraph(s)."""
    tag = element.tag
    align = _block_align(element, default_align)

    if tag == "p":
        para = _new_para(doc, align)
        _parse_inline(element, para, italic=default_italic, font_size=default_size)

    elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        sizes = {1: 16, 2: 14, 3: 13, 4: 12, 5: 12, 6: 11}
        para = _new_para(doc, align or WD_ALIGN_PARAGRAPH.LEFT)
        _parse_inline(element, para, bold=True, font_size=float(sizes.get(level, 13)))

    elif tag in ("ul", "ol"):
        counter = 1
        for li in element:
            if getattr(li, "tag", None) != "li":
                continue
            para = _new_para(doc, default_align)
            # Hanging indent: number/bullet at 0.5cm, text wraps at 1.0cm
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.first_line_indent = Cm(-0.5)

            marker_text = "•  " if tag == "ul" else f"{counter}. "
            counter += 1
            marker_run = para.add_run(marker_text)
            marker_run.font.name = "Times New Roman"
            marker_run.font.size = Pt(default_size)
            marker_run.font.italic = default_italic

            # TipTap wraps list-item content in <p>; fall back to li itself
            p_child = next((c for c in li if getattr(c, "tag", None) == "p"), None)
            _parse_inline(
                p_child if p_child is not None else li,
                para,
                italic=default_italic,
                font_size=default_size,
            )

    else:
        # Recurse into unknown block containers (div, blockquote, …)
        for child in element:
            if getattr(child, "tag", None):
                _process_block(doc, child, default_align, default_size, default_italic)


def _html_section(
    doc: Document,
    html: str,
    default_align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    default_size: float = 14.0,
    default_italic: bool = False,
) -> bool:
    """Parse TipTap HTML and append paragraphs to *doc*. Returns True if any content added."""
    if not html or html.strip() in ("", "<p></p>"):
        return False
    try:
        root = lxml_html.fragment_fromstring(html, create_parent="div")
    except Exception:
        text = re.sub(r"<[^>]+>", " ", html).strip()
        if text:
            para = _new_para(doc, default_align)
            run = para.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(default_size)
            run.font.italic = default_italic
            return True
        return False

    added = False
    for child in root:
        if getattr(child, "tag", None):
            _process_block(doc, child, default_align, default_size, default_italic)
            added = True
    return added


# ── Shared helpers ────────────────────────────────────────────────────────────

def _set_font(run, name: str = "Times New Roman", size_pt: float = 13,
              bold: bool = False, italic: bool = False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic


def _add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
              size_pt: float = 13,
              align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_font(run, size_pt=size_pt, bold=bold, italic=italic)
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    return para


def _remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tcBorders.append(border)
    tcPr.append(tcBorders)


# ── Main document builder ─────────────────────────────────────────────────────

def _build_docx_impl(data: dict[str, Any]) -> bytes:
    doc = Document()

    # Page margins — NĐ30: top 25mm, bottom 25mm, left 30mm, right 20mm
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    loai          = data.get("loaiVanBan", "QĐ")
    cq_chu_quan   = data.get("coQuanChuQuan", "")
    cq_ban_hanh   = data.get("coQuanBanHanh", "")
    so_ky_hieu    = data.get("soKyHieu", "")
    dia_danh      = data.get("diaDanh", "")
    ngay_thang    = data.get("ngayThang", "")
    trich_yeu     = data.get("trichYeu", "")
    kinh_gui      = data.get("kinhGui", "")
    can_cu_html   = data.get("canCu", "")
    noi_dung_html = data.get("noiDung", "")
    quyen_han     = data.get("quyenHanKy", "")
    chuc_vu       = data.get("chucVuKy", "")
    ho_ten        = data.get("hoTenKy", "")
    noi_nhan      = data.get("noiNhan", []) or []
    chuc_danh_tap_the = data.get("chucDanhTapThe", "")

    _ABBR_MAP = {
        "NQ": "NGHỊ QUYẾT",  "QĐ": "QUYẾT ĐỊNH", "CT": "CHỈ THỊ",
        "CV": "CÔNG VĂN",    "TB": "THÔNG BÁO",   "HD": "HƯỚNG DẪN",
        "BC": "BÁO CÁO",     "TTr": "TỜ TRÌNH",   "KH": "KẾ HOẠCH",
        "BB": "BIÊN BẢN",    "GM": "GIẤY MỜI",
    }
    ten_loai = _ABBR_MAP.get(loai, loai.upper())

    # ── Header table (2 columns, no borders) ──────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = doc.styles["Normal Table"]
    left_cell  = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)
    _remove_cell_borders(left_cell)
    _remove_cell_borders(right_cell)

    # Left column: cơ quan chủ quản + ban hành + số/KH
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cq_chu_quan:
        r = lp.add_run(cq_chu_quan.upper())
        _set_font(r, size_pt=12)
        lp = left_cell.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = lp.add_run(cq_ban_hanh.upper() if cq_ban_hanh else "")
    _set_font(r, size_pt=12, bold=True)

    lp2 = left_cell.add_paragraph("─" * 18)
    lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(lp2.runs[0] if lp2.runs else lp2.add_run(""), size_pt=11)

    lp3 = left_cell.add_paragraph()
    lp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = lp3.add_run(f"Số: {so_ky_hieu}")
    _set_font(r, size_pt=13, italic=True)

    # Right column: quốc hiệu + địa danh/ngày
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rp.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    _set_font(r, size_pt=12, bold=True)

    rp2 = right_cell.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rp2.add_run("Độc lập - Tự do - Hạnh phúc")
    _set_font(r, size_pt=13, bold=True)

    rp3 = right_cell.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rp3.add_run("─" * 28)
    _set_font(r, size_pt=11)

    dia_ngay = ", ".join(filter(None, [dia_danh, ngay_thang]))
    rp4 = right_cell.add_paragraph()
    rp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rp4.add_run(dia_ngay)
    _set_font(r, size_pt=13, italic=True)

    doc.add_paragraph()  # spacing

    # ── Tên loại VB + trích yếu ───────────────────────────────────────────────
    _add_para(doc, ten_loai, bold=True, size_pt=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    if trich_yeu:
        _add_para(doc, trich_yeu, bold=True, size_pt=13, align=WD_ALIGN_PARAGRAPH.CENTER)

    div_p = _add_para(doc, "─" * 38, size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    div_p.paragraph_format.space_after = Pt(6)

    # ── Kính gửi ──────────────────────────────────────────────────────────────
    if kinh_gui:
        _add_para(doc, f"Kính gửi: {kinh_gui}", size_pt=14,
                  align=WD_ALIGN_PARAGRAPH.LEFT)
        doc.add_paragraph()

    # ── Căn cứ (italic, 14pt, justify) ───────────────────────────────────────
    has_can_cu = _html_section(
        doc, can_cu_html,
        default_align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        default_size=14.0,
        default_italic=True,
    )
    if has_can_cu:
        doc.add_paragraph()

    # ── Nội dung (14pt, justify) ──────────────────────────────────────────────
    _html_section(
        doc, noi_dung_html,
        default_align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        default_size=14.0,
        default_italic=False,
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Footer table: Nơi nhận | Chữ ký ──────────────────────────────────────
    ftr = doc.add_table(rows=1, cols=2)
    ftr.style = doc.styles["Normal Table"]
    fl = ftr.cell(0, 0)
    fr = ftr.cell(0, 1)
    _remove_cell_borders(fl)
    _remove_cell_borders(fr)

    lp = fl.paragraphs[0]
    r = lp.add_run("Nơi nhận:")
    _set_font(r, bold=True, italic=True, size_pt=12)
    for item in noi_nhan:
        p = fl.add_paragraph(item if item.startswith("-") else f"- {item}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            _set_font(run, size_pt=11)

    rp = fr.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ky_header = " ".join(filter(None, [quyen_han, chuc_danh_tap_the]))
    r = rp.add_run(ky_header.upper() if ky_header else "")
    _set_font(r, bold=True, size_pt=13)

    if chuc_vu:
        p_cv = fr.add_paragraph()
        p_cv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_cv.add_run(chuc_vu.upper())
        _set_font(r, bold=True, size_pt=13)

    for _ in range(4):
        p_gap = fr.add_paragraph()
        p_gap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if ho_ten:
        p_ht = fr.add_paragraph()
        p_ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_ht.add_run(ho_ten)
        _set_font(r, bold=True, size_pt=13)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx(data: dict[str, Any]) -> bytes:
    try:
        return _build_docx_impl(data)
    except Exception as exc:
        logger.error("DOCX generation failed: %s", exc, exc_info=True)
        raise


async def generate_docx(data: dict[str, Any]) -> bytes:
    """Render Nd30Data dict → DOCX bytes (runs python-docx in a thread)."""
    return await asyncio.to_thread(_build_docx, data)
