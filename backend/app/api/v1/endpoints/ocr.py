"""
OCR endpoints — async job-based (DB + Redis) + stateless export.

POST /extract   → create OcrJob (202), run OCR in background
POST /export    → stateless text→DOCX/PDF (unchanged)
GET  /jobs      → list jobs for current user
GET  /status/{job_id} → lightweight poll (no text payload)
GET  /{job_id}  → full job detail

Route ordering: /jobs and /status/{id} MUST come before /{id}.
"""
import asyncio
import base64
import io
import json
import logging
import re
import uuid
from urllib.parse import quote

from fastapi import (
    APIRouter, BackgroundTasks, Body, Depends,
    File, HTTPException, Query, Response, UploadFile, status,
)
from sqlalchemy import func as sql_func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_user
from app.core.database import AsyncSessionLocal, get_db
from app.core.redis import get_redis
from app.core.storage import delete_file, download_file, get_bucket_name, upload_file
from app.models.ocr_job import OcrJob
from app.models.user import User
from app.schemas.ocr_job import OcrJobListResponse, OcrJobResponse, OcrJobStatusResponse

logger = logging.getLogger(__name__)
router = APIRouter()

_ALLOWED_TYPES = {
    "application/pdf",
    "image/jpeg",
    "image/jpg",
    "image/png",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_MAX_BYTES = 20 * 1024 * 1024  # 20 MB
_REDIS_TTL = 3600               # 1 hour
_OCR_BUCKET = get_bucket_name()

_REVIEW_SYSTEM_PROMPT = """\
Bạn là chuyên gia rà soát văn bản hành chính Việt Nam.
Nhiệm vụ: Phân tích và chỉnh sửa văn bản theo các tiêu chí:
1. Chính tả và typo (sửa lỗi viết sai)
2. Thể thức văn bản theo Nghị định 30/2020/NĐ-CP
3. Văn phong hành chính (trang trọng, rõ ràng, súc tích)
4. Dấu câu và cách trình bày
5. Thuật ngữ pháp lý đúng chuẩn

Trả về JSON với format:
{
  "reviewed_text": "toàn bộ văn bản đã chỉnh sửa",
  "changes": [
    {
      "type": "chinh_ta|the_thuc|van_phong|dau_cau|thuat_ngu",
      "original": "đoạn gốc",
      "revised": "đoạn đã sửa",
      "reason": "lý do chỉnh sửa"
    }
  ],
  "summary": "tóm tắt các điểm đã chỉnh sửa"
}
Chỉ trả về JSON, không thêm gì khác.\
"""


# ── LLM text formatter ───────────────────────────────────────────────────────

def _basic_format(text: str) -> str:
    """Fallback: normalize whitespace when LLM is unavailable."""
    import re
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()


async def _format_ocr_text(raw_text: str, filename: str) -> str:
    """Use LLM to reformat raw OCR text into clean, readable Vietnamese text."""
    # Cap at 8000 chars to stay well within Groq context / avoid timeout
    truncated = raw_text[:8000] if len(raw_text) > 8000 else raw_text

    system_prompt = (
        "Bạn là trợ lý chuyên xử lý văn bản tiếng Việt được trích xuất từ OCR.\n"
        "Nhiệm vụ: nhận text OCR thô (có thể bị lỗi định dạng, dính chữ, mất xuống dòng) "
        "và tái cấu trúc lại thành văn bản có định dạng đẹp, dễ đọc.\n\n"
        "Nguyên tắc:\n"
        "- Giữ NGUYÊN nội dung, KHÔNG thêm hoặc bịa thông tin\n"
        "- Tách đúng các đoạn văn, tiêu đề, danh sách\n"
        "- Nếu là văn bản hành chính (công văn, quyết định...): giữ cấu trúc quốc hiệu, "
        "số ký hiệu, trích yếu, nội dung, chữ ký\n"
        "- Nếu là biểu mẫu (tờ khai, hợp đồng...): giữ cấu trúc các mục, trường thông tin\n"
        "- Nếu là văn bản thông thường: tách paragraph hợp lý\n"
        "- Dùng dấu xuống dòng \\n để phân tách đoạn\n"
        "- KHÔNG thêm markdown (không dùng **, ##, -)\n"
        "- Chỉ trả về văn bản đã tái cấu trúc, KHÔNG giải thích gì thêm"
    )
    user_prompt = f"Tên file: {filename}\n\nText OCR cần tái cấu trúc:\n{truncated}"

    try:
        from app.services.llm_service import llm_service
        result = await llm_service.chat(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.1,
            max_tokens=4000,
        )
        return result.strip() if result else raw_text
    except Exception as exc:
        logger.warning("[ocr_format] LLM failed: %s — falling back to basic format", exc)
        return _basic_format(raw_text)


# ── PDF generator from OCR text ──────────────────────────────────────────────

async def _create_pdf_from_text(text: str, filename: str, job_id: str) -> str | None:
    """
    Tạo PDF từ formatted_text, upload lên R2.
    Best-effort — lỗi chỉ log warning, không fail cả job.
    Returns R2 object key hoặc None nếu lỗi.
    """
    try:
        import html as _html_esc
        from app.services.pdf_service import _write_pdf, _ensure_fonts, _build_css

        font = _ensure_fonts()
        css = _build_css(font)

        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        if not paragraphs:
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

        title_h = _html_esc.escape(filename)
        body_html = "".join(
            f"<p>{_html_esc.escape(p).replace(chr(10), '<br/>')}</p>"
            for p in paragraphs
        )
        html_str = (
            "<!DOCTYPE html>\n<html lang='vi'>\n<head>\n"
            "  <meta charset='UTF-8'>\n"
            f"  <style>{css}\n"
            f"    h1 {{ font-size: 15pt; font-weight: bold; text-align: center;"
            f" margin-bottom: 4pt; }}\n"
            f"    p {{ font-family: '{font}', serif; font-size: 13pt;"
            f" line-height: 1.6; margin: 0 0 0.6em 0; }}\n"
            "  </style>\n</head>\n<body>\n"
            f"  <h1>{title_h}</h1>\n"
            f"  {body_html}\n"
            "</body>\n</html>"
        )

        pdf_bytes = await asyncio.to_thread(_write_pdf, html_str)
        if not pdf_bytes:
            logger.warning("[ocr_pdf] %s _write_pdf returned empty", job_id)
            return None

        safe_name = re.sub(r"[^\w\-.]", "_", filename.rsplit(".", 1)[0])
        r2_key = f"ocr-jobs/{job_id}/{safe_name}_ocr.pdf"
        await upload_file(pdf_bytes, r2_key, "application/pdf", bucket_name=_OCR_BUCKET)
        logger.info("[ocr_pdf] %s uploaded %d bytes → %s", job_id, len(pdf_bytes), r2_key)
        return r2_key

    except Exception as exc:
        logger.warning("[ocr_pdf] %s failed to create PDF: %s", job_id, exc)
        return None


# ── OCR PDF with per-page progress ───────────────────────────────────────────

async def _ocr_pdf_with_progress(content: bytes, job_id: str) -> str:
    """OCR từng trang PDF, ghi Redis progress sau mỗi trang."""
    import json
    import sys

    redis = await get_redis()

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            total_pages = len(pdf.pages)
    except Exception:
        total_pages = 1

    await redis.set(
        f"ocr_progress:{job_id}",
        json.dumps({"current_page": 0, "total_pages": total_pages}),
        ex=3600,
    )

    def convert_pages():
        import pdf2image
        return pdf2image.convert_from_bytes(content, dpi=300)

    images = await asyncio.to_thread(convert_pages)
    total_pages = len(images)

    all_text: list[str] = []
    for i, image in enumerate(images):
        def ocr_page(img):
            import pytesseract
            if sys.platform == "win32":
                pytesseract.pytesseract.tesseract_cmd = (
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                )
            try:
                return pytesseract.image_to_string(img, lang="vie")
            except Exception:
                return pytesseract.image_to_string(img, lang="eng")

        page_text = await asyncio.to_thread(ocr_page, image)
        all_text.append(page_text)

        await redis.set(
            f"ocr_progress:{job_id}",
            json.dumps({"current_page": i + 1, "total_pages": total_pages}),
            ex=3600,
        )

    await redis.delete(f"ocr_progress:{job_id}")
    return "\n\n".join(all_text)


# ── PDF type detection ────────────────────────────────────────────────────────

async def _detect_pdf_type(content: bytes) -> str:
    """
    Detect PDF có text layer hay scan.
    Returns: "text_pdf" hoặc "scanned_pdf"
    """
    import pdfplumber
    import io
    try:
        def extract():
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = ""
                for page in pdf.pages:
                    text += (page.extract_text() or "")
                    if len(text) > 50:
                        return "text_pdf"
            return "scanned_pdf" if len(text.strip()) < 50 else "text_pdf"
        return await asyncio.to_thread(extract)
    except Exception:
        return "scanned_pdf"  # fallback: treat as scan


# ── Background processor ──────────────────────────────────────────────────────

async def _process_ocr_job(job_id: str) -> None:
    """
    Background task: OCR the cached file bytes and persist result.

    Three separate DB sessions to avoid using a rolled-back session for the
    error-state update (same pattern as pipeline_service.py).
    Redis stores file bytes as base64 because the client uses decode_responses=True.
    """

    # ── Phase 1: mark as processing, capture filename ─────────────────────
    filename = ""
    try:
        async with AsyncSessionLocal() as db:
            result = await db.execute(select(OcrJob).where(OcrJob.id == job_id))
            job = result.scalar_one_or_none()
            if not job:
                logger.warning("[ocr_job] %s not found in DB — skipping", job_id)
                return
            filename = job.filename or ""
            job.status = "processing"
            await db.commit()
    except Exception as exc:
        logger.exception("[ocr_job] %s failed to mark processing: %s", job_id, exc)
        return

    # ── Phase 2: OCR work + LLM formatting ───────────────────────────────
    extracted_text: str = ""
    formatted: str = ""
    page_count: int = 1
    error_msg: str | None = None
    pdf_path: str | None = None

    try:
        redis = await get_redis()
        raw = await redis.get(f"ocr_file:{job_id}")
        if not raw:
            raise RuntimeError("File không còn trong cache")

        # decode_responses=True → raw is str; base64.b64decode accepts str directly
        content: bytes = base64.b64decode(raw)
        is_pdf = filename.lower().endswith(".pdf")

        if is_pdf:
            import pdfplumber

            # Extract text to detect scanned vs text PDF (< 50 chars → scanned)
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    page_count = len(pdf.pages)
                    pdf_text = "".join(
                        (p.extract_text() or "") for p in pdf.pages
                    ).strip()
            except Exception:
                page_count = 1
                pdf_text = ""

            if len(pdf_text) < 50:
                # Scanned PDF → OCR with per-page progress tracking
                extracted_text = await _ocr_pdf_with_progress(content, job_id)
            else:
                # Text PDF → fast path
                from app.services.pipeline_service import _extract_pdf
                extracted_text = await asyncio.to_thread(_extract_pdf, content)

        else:
            def _ocr_image(data: bytes) -> str:
                import sys
                import pytesseract
                from PIL import Image

                if sys.platform == "win32":
                    pytesseract.pytesseract.tesseract_cmd = (
                        r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                    )
                img = Image.open(io.BytesIO(data))
                try:
                    return pytesseract.image_to_string(img, lang="vie")
                except Exception:
                    return pytesseract.image_to_string(img, lang="eng")

            extracted_text = await asyncio.to_thread(_ocr_image, content)
            page_count = 1

        # Cleanup Redis (best-effort — TTL covers it if this fails)
        await redis.delete(f"ocr_file:{job_id}")
        logger.info(
            "[ocr_job] %s OCR done: %d chars, %d pages", job_id, len(extracted_text), page_count
        )

        # LLM formatting step — non-fatal, falls back to _basic_format on error
        formatted = await _format_ocr_text(extracted_text, filename)
        logger.info("[ocr_job] %s formatted: %d chars", job_id, len(formatted))

        # Generate PDF for unified viewer (best-effort)
        pdf_path = await _create_pdf_from_text(formatted or extracted_text, filename, job_id)

    except Exception as exc:
        logger.exception("[ocr_job] %s OCR failed: %s", job_id, exc)
        error_msg = str(exc)[:500]

    # ── Phase 3: persist result (fresh session — always runs) ─────────────
    try:
        async with AsyncSessionLocal() as db:
            result = await db.execute(select(OcrJob).where(OcrJob.id == job_id))
            job = result.scalar_one_or_none()
            if not job:
                return
            if error_msg:
                job.status = "error"
                job.error_msg = error_msg
            else:
                job.status = "done"
                job.text = extracted_text
                job.formatted_text = formatted
                job.page_count = page_count
                job.char_count = len(extracted_text)
                job.file_path = pdf_path
            await db.commit()
    except Exception as exc:
        logger.exception("[ocr_job] %s failed to persist result: %s", job_id, exc)


# ── Background processor — text PDF ──────────────────────────────────────────

async def _process_text_pdf(job_id: str, content: bytes) -> None:
    """Background task cho PDF văn bản — extract text + LLM format."""

    # Phase 1: update status processing
    filename = ""
    try:
        async with AsyncSessionLocal() as db:
            job = await db.get(OcrJob, job_id)
            if not job:
                return
            filename = job.filename or ""
            job.status = "processing"
            await db.commit()
    except Exception as exc:
        logger.exception("[text_pdf] %s failed to mark processing: %s", job_id, exc)
        return

    # Phase 2: extract text + count pages + LLM format
    try:
        import pdfplumber

        def extract_all():
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                return "\n\n".join(pages), len(pdf.pages)

        extracted_text, page_count = await asyncio.to_thread(extract_all)
        formatted = await _format_ocr_text(extracted_text, filename)
        logger.info("[text_pdf] %s done: %d chars, %d pages", job_id, len(extracted_text), page_count)

    except Exception as exc:
        logger.exception("[text_pdf] %s failed: %s", job_id, exc)
        async with AsyncSessionLocal() as db:
            job = await db.get(OcrJob, job_id)
            if job:
                job.status = "error"
                job.error_msg = str(exc)[:500]
                await db.commit()
        return

    # Phase 3: save done
    try:
        async with AsyncSessionLocal() as db:
            job = await db.get(OcrJob, job_id)
            if job:
                job.status = "done"
                job.text = extracted_text
                job.formatted_text = formatted
                job.page_count = page_count
                job.char_count = len(extracted_text)
                await db.commit()
    except Exception as exc:
        logger.exception("[text_pdf] %s failed to persist result: %s", job_id, exc)


# ── Background processor — DOCX ──────────────────────────────────────────────

async def _process_text_docx(job_id: str, content: bytes) -> None:
    """Background task cho DOCX — extract text + LLM format."""

    # Phase 1: mark processing
    filename = ""
    try:
        async with AsyncSessionLocal() as db:
            job = await db.get(OcrJob, job_id)
            if not job:
                return
            filename = job.filename or ""
            job.status = "processing"
            await db.commit()
    except Exception as exc:
        logger.exception("[text_docx] %s failed to mark processing: %s", job_id, exc)
        return

    # Phase 2: extract text + LLM format
    try:
        from docx import Document as DocxDocument

        def extract_docx():
            doc = DocxDocument(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)

        extracted_text = await asyncio.to_thread(extract_docx)
        page_count = 1  # DOCX không có concept trang rõ ràng
        formatted = await _format_ocr_text(extracted_text, filename)
        logger.info("[text_docx] %s done: %d chars", job_id, len(extracted_text))

        # Generate PDF for unified viewer (best-effort)
        pdf_path = await _create_pdf_from_text(formatted or extracted_text, filename, job_id)

    except Exception as exc:
        logger.exception("[text_docx] %s failed: %s", job_id, exc)
        async with AsyncSessionLocal() as db:
            job = await db.get(OcrJob, job_id)
            if job:
                job.status = "error"
                job.error_msg = str(exc)[:500]
                await db.commit()
        return

    # Phase 3: save done
    try:
        async with AsyncSessionLocal() as db:
            job = await db.get(OcrJob, job_id)
            if job:
                job.status = "done"
                job.text = extracted_text
                job.formatted_text = formatted
                job.page_count = page_count
                job.char_count = len(extracted_text)
                job.file_path = pdf_path
                await db.commit()
    except Exception as exc:
        logger.exception("[text_docx] %s failed to persist result: %s", job_id, exc)


# ── POST /extract ─────────────────────────────────────────────────────────────

@router.post("/extract", response_model=OcrJobResponse, status_code=status.HTTP_202_ACCEPTED)
async def ocr_extract(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Start async OCR job. Returns OcrJob (status=pending) immediately."""
    content_type = (file.content_type or "").lower()
    if content_type not in _ALLOWED_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Chỉ hỗ trợ PDF, JPG, PNG, DOCX",
        )

    content = await file.read()
    if len(content) > _MAX_BYTES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File quá lớn, tối đa 20MB",
        )

    # Detect file type: text_pdf / scanned_pdf / image / text_docx
    _docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    is_docx = content_type == _docx_mime or (file.filename or "").lower().endswith(".docx")
    is_pdf = content_type == "application/pdf"

    if is_docx:
        file_type = "text_docx"
    elif is_pdf:
        file_type = await _detect_pdf_type(content)
    else:
        file_type = "image"

    job_id = str(uuid.uuid4())
    filename = file.filename or "unknown"
    file_path: str | None = None

    if file_type == "text_pdf":
        # Upload original file to R2 immediately for later download
        r2_key = f"ocr-jobs/{job_id}/{filename}"
        await upload_file(content, r2_key, "application/pdf", bucket_name=_OCR_BUCKET)
        file_path = r2_key

    job = OcrJob(
        id=job_id,
        user_id=str(current_user.id),
        filename=filename,
        status="pending",
        file_type=file_type,
        file_path=file_path,
    )
    db.add(job)
    await db.flush()
    await db.commit()

    if file_type == "text_pdf":
        background_tasks.add_task(_process_text_pdf, job.id, content)
    elif file_type == "text_docx":
        background_tasks.add_task(_process_text_docx, job.id, content)
    else:
        # scanned_pdf / image — cache in Redis for _process_ocr_job
        redis = await get_redis()
        await redis.set(
            f"ocr_file:{job.id}",
            base64.b64encode(content).decode(),
            ex=_REDIS_TTL,
        )
        background_tasks.add_task(_process_ocr_job, job.id)

    logger.info(
        "[ocr] queued job=%s file=%s type=%s bytes=%d",
        job.id, job.filename, file_type, len(content),
    )
    return job


# ── POST /export ──────────────────────────────────────────────────────────────

@router.post("/export")
async def ocr_export(
    text: str = Body(...),
    filename: str = Body("van-ban-ocr"),
    format: str = Body("docx"),
    current_user: User = Depends(get_current_user),
):
    """Convert OCR text to DOCX or PDF (stateless — nothing saved)."""
    if format not in ("docx", "pdf"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="format phải là 'docx' hoặc 'pdf'",
        )

    safe_name = re.sub(r"[^\w\-.]", "_", filename)[:120]

    # Split into paragraphs (prefer double-newline, fall back to single)
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

    if format == "docx":
        from docx import Document as DocxDocument
        from docx.shared import Pt

        doc = DocxDocument()
        doc.add_heading(filename, level=1)
        for para_text in paragraphs:
            para = doc.add_paragraph()
            run = para.add_run(para_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
        buf = io.BytesIO()
        doc.save(buf)
        encoded = quote(f"{safe_name}.docx", safe="")
        return Response(
            content=buf.getvalue(),
            media_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": (
                    f"attachment; filename=\"vanban.docx\"; "
                    f"filename*=UTF-8''{encoded}"
                )
            },
        )

    # format == "pdf"
    import html as _html_esc
    from app.services.pdf_service import _build_css, _ensure_fonts, _write_pdf

    font = _ensure_fonts()
    css = _build_css(font)
    title_h = _html_esc.escape(filename)
    body_html = "".join(
        f"<p>{_html_esc.escape(p).replace(chr(10), '<br/>')}</p>"
        for p in paragraphs
    )
    html_str = (
        "<!DOCTYPE html>\n<html lang='vi'>\n<head>\n"
        "  <meta charset='UTF-8'>\n"
        f"  <style>{css}\n"
        f"    h1 {{ font-size: 15pt; font-weight: bold; text-align: center;"
        f" margin-bottom: 4pt; }}\n"
        f"    p {{ font-family: '{font}', serif; font-size: 13pt;"
        f" line-height: 1.6; margin: 0 0 0.6em 0; }}\n"
        "  </style>\n</head>\n<body>\n"
        f"  <h1>{title_h}</h1>\n"
        f"  {body_html}\n"
        "</body>\n</html>"
    )
    pdf_bytes = await asyncio.to_thread(_write_pdf, html_str)
    encoded = quote(f"{safe_name}.pdf", safe="")
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": (
                f"attachment; filename=\"vanban.pdf\"; "
                f"filename*=UTF-8''{encoded}"
            )
        },
    )


# ── GET /jobs ─────────────────────────────────────────────────────────────────

@router.get("/jobs", response_model=OcrJobListResponse)
async def list_ocr_jobs(
    skip: int = 0,
    limit: int = 20,
    status_filter: str | None = Query(None, alias="status"),
    file_type: str | None = Query(None),
    sort_by: str = Query("created_at"),
    sort_order: str = Query("desc"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List OCR jobs for the current user with optional filter and sort."""
    base = select(OcrJob).where(OcrJob.user_id == str(current_user.id))

    if status_filter:
        base = base.where(OcrJob.status == status_filter)
    if file_type:
        base = base.where(OcrJob.file_type == file_type)

    total_result = await db.execute(
        select(sql_func.count()).select_from(base.subquery())
    )
    total = total_result.scalar() or 0

    _sort_cols = {
        "created_at": OcrJob.created_at,
        "filename":   OcrJob.filename,
        "page_count": OcrJob.page_count,
    }
    sort_col = _sort_cols.get(sort_by, OcrJob.created_at)
    sort_expr = sort_col.asc() if sort_order == "asc" else sort_col.desc()

    items_result = await db.execute(
        base.order_by(sort_expr).offset(skip).limit(limit)
    )
    items = items_result.scalars().all()

    return OcrJobListResponse(items=list(items), total=total)


# ── GET /status/{job_id} ──────────────────────────────────────────────────────

@router.get("/status/{job_id}", response_model=OcrJobStatusResponse)
async def get_ocr_status(
    job_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Lightweight status poll — no text payload in response."""
    result = await db.execute(
        select(OcrJob).where(
            OcrJob.id == job_id,
            OcrJob.user_id == str(current_user.id),
        )
    )
    job = result.scalar_one_or_none()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Job not found")
    return job


# ── GET /progress/{job_id} ───────────────────────────────────────────────────

@router.get("/progress/{job_id}")
async def get_ocr_progress(
    job_id: str,
    current_user: User = Depends(get_current_user),
):
    """Trả về tiến độ OCR theo trang từ Redis."""
    import json
    redis = await get_redis()
    raw = await redis.get(f"ocr_progress:{job_id}")
    if not raw:
        return {"job_id": job_id, "current_page": 0, "total_pages": 0, "percent": 0}

    data = json.loads(raw)
    current = data.get("current_page", 0)
    total = data.get("total_pages", 0)
    percent = int((current / total) * 100) if total > 0 else 0
    return {
        "job_id": job_id,
        "current_page": current,
        "total_pages": total,
        "percent": percent,
    }


# ── GET /{job_id}/export/docx ────────────────────────────────────────────────

@router.get("/{job_id}/export/docx")
async def export_ocr_as_docx(
    job_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Convert file gốc PDF sang DOCX dùng pdf2docx."""
    try:
        from pdf2docx import Converter  # noqa: F401
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Tính năng chuyển đổi Word tạm thời không khả dụng",
        )

    result = await db.execute(
        select(OcrJob).where(
            OcrJob.id == job_id,
            OcrJob.user_id == str(current_user.id),
        )
    )
    job = result.scalar_one_or_none()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Job not found")

    if job.file_type != "text_pdf" or not job.file_path:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Chỉ hỗ trợ xuất Word cho PDF văn bản",
        )

    # Download PDF gốc từ R2
    pdf_bytes = await asyncio.to_thread(download_file, job.file_path, _OCR_BUCKET)

    # Convert PDF → DOCX (blocking I/O — chạy trong thread)
    def convert(pdf_data: bytes) -> bytes:
        import os
        import tempfile
        from pdf2docx import Converter

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_tmp:
            pdf_tmp.write(pdf_data)
            pdf_path = pdf_tmp.name

        docx_path = pdf_path.replace(".pdf", ".docx")
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            with open(docx_path, "rb") as f:
                return f.read()
        finally:
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
            if os.path.exists(docx_path):
                os.unlink(docx_path)

    try:
        docx_bytes = await asyncio.to_thread(convert, pdf_bytes)
    except Exception as exc:
        logger.error("[export_docx] job %s convert failed: %s", job_id, exc)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Không thể convert PDF sang Word",
        )

    safe_name = re.sub(r"[^\w\-.]", "_", job.filename.rsplit(".", 1)[0])
    encoded = quote(f"{safe_name}.docx", safe="")
    return Response(
        content=docx_bytes,
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": (
                f"attachment; filename=\"{safe_name}.docx\"; "
                f"filename*=UTF-8''{encoded}"
            )
        },
    )


# ── GET /{job_id}/download ────────────────────────────────────────────────────

@router.get("/{job_id}/download")
async def download_ocr_file(
    job_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Stream PDF từ R2 — text_pdf: file gốc; scan/image/docx: PDF được tạo từ OCR text."""
    result = await db.execute(
        select(OcrJob).where(
            OcrJob.id == job_id,
            OcrJob.user_id == str(current_user.id),
        )
    )
    job = result.scalar_one_or_none()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Job not found")

    if not job.file_path:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File chưa sẵn sàng",
        )

    data = await asyncio.to_thread(download_file, job.file_path, _OCR_BUCKET)
    encoded = quote(job.filename, safe="")
    return Response(
        content=data,
        media_type="application/pdf",
        headers={
            "Content-Disposition": (
                f"attachment; filename=\"{job.filename}\"; "
                f"filename*=UTF-8''{encoded}"
            )
        },
    )


# ── POST /{job_id}/review ────────────────────────────────────────────────────

@router.post("/{job_id}/review")
async def review_ocr_job(
    job_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Gọi LLM rà soát chính tả/thể thức văn bản OCR. Kết quả cache Redis 24h."""
    result = await db.execute(
        select(OcrJob).where(
            OcrJob.id == job_id,
            OcrJob.user_id == str(current_user.id),
        )
    )
    job = result.scalar_one_or_none()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Job not found")
    if job.status != "done":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Job chưa hoàn thành",
        )

    text = job.formatted_text or job.text
    if not text or not text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Không có nội dung để rà soát",
        )

    try:
        from app.services.llm_service import llm_service
        if not llm_service._base_url:
            raise HTTPException(status_code=503, detail="LLM service không khả dụng")

        raw = await llm_service.chat(
            messages=[
                {"role": "system", "content": _REVIEW_SYSTEM_PROMPT},
                {"role": "user", "content": text[:6000]},
            ],
            temperature=0.1,
            max_tokens=4000,
            json_mode=True,
        )
        review_data = json.loads(raw)
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("[review_ocr] job=%s error: %s", job_id, exc)
        raise HTTPException(status_code=503, detail=f"LLM error: {exc}")

    redis = await get_redis()
    await redis.set(
        f"review:{job_id}",
        json.dumps(review_data, ensure_ascii=False),
        ex=86400,
    )

    return {
        "job_id": job_id,
        "reviewed_text": review_data.get("reviewed_text", ""),
        "changes": review_data.get("changes", []),
        "summary": review_data.get("summary", ""),
    }


# ── GET /{job_id} ─────────────────────────────────────────────────────────────

@router.get("/{job_id}", response_model=OcrJobResponse)
async def get_ocr_job(
    job_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Full job detail including extracted text."""
    result = await db.execute(
        select(OcrJob).where(
            OcrJob.id == job_id,
            OcrJob.user_id == str(current_user.id),
        )
    )
    job = result.scalar_one_or_none()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Job not found")
    return job


# ── DELETE /{job_id} ──────────────────────────────────────────────────────────

@router.delete("/{job_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_ocr_job(
    job_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Xóa OCR job và file liên quan trên storage."""
    result = await db.execute(
        select(OcrJob).where(
            OcrJob.id == job_id,
            OcrJob.user_id == str(current_user.id),
        )
    )
    job = result.scalar_one_or_none()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Job not found")

    if job.file_path:
        try:
            await asyncio.to_thread(delete_file, job.file_path, _OCR_BUCKET)
        except Exception as exc:
            logger.warning("[ocr_delete] %s failed to delete file %s: %s", job_id, job.file_path, exc)

    await db.delete(job)
    await db.commit()
