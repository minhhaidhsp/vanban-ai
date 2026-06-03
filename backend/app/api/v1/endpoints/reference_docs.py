from fastapi import APIRouter, BackgroundTasks, Depends, Form, HTTPException, Response, status, UploadFile, File, Query
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func as sql_func, or_, text
from typing import List
from app.core.config import get_settings
from app.core.database import get_db
from app.core.redis import get_redis
from app.core.storage import (
    get_storage_client as get_minio_client,
    ensure_bucket_exists, get_file_url, get_bucket_name,
    upload_file_data, download_file_data, delete_file_data,
)
from app.api.deps import get_current_user
from app.models.user import User
from app.models.reference_document import ReferenceDocument
from app.models.reference_doc_chunk import ReferenceDocChunk
from app.services.reference_pipeline_service import process_reference_background
from app.schemas.reference_document import (
    RefDocCreate, RefDocUpdate, RefDocResponse,
    RefDocListResponse, RefDocSearchResponse,
    RefDocChunkSearchItem, RefDocChunkSearchResponse,
    RefDocFTSItem, RefDocFTSResponse,
    MetadataPreviewResponse, MetadataConfirmRequest, MetadataConfidence,
    ChunkItem, RefDocContentResponse,
)
import asyncio
import io
import json
import logging
import re
import uuid
from pydantic import BaseModel

logger = logging.getLogger(__name__)

router = APIRouter()

REF_DOCS_BUCKET = get_bucket_name()

_VALID_HIEU_LUC = {"con_hieu_luc", "het_hieu_luc", "mot_phan", "chua"}


class HieuLucUpdate(BaseModel):
    hieu_luc: str
    ghi_chu: str | None = None


# ── Helpers ──────────────────────────────────────────────────────────────────

def _add_url(
    doc: ReferenceDocument,
    score: float | None = None,
    chunk_count: int | None = None,
) -> RefDocResponse:
    resp = RefDocResponse.model_validate(doc)
    if doc.file_path:
        try:
            resp.download_url = get_file_url(doc.file_path, bucket_name=REF_DOCS_BUCKET)
        except Exception:
            pass
    if score is not None:
        resp.score = score
    if chunk_count is not None:
        resp.chunk_count = chunk_count
    return resp


# ── List ─────────────────────────────────────────────────────────────────────

_SORT_COLUMNS = {
    "created_at": ReferenceDocument.created_at,
    "title":      ReferenceDocument.title,
    "loai_van_ban": ReferenceDocument.loai_van_ban,
    "ngay_ban_hanh": ReferenceDocument.ngay_ban_hanh,
}


@router.get("/", response_model=RefDocListResponse)
async def list_ref_docs(
    skip: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
    loai: str | None = Query(None),
    hieu_luc: str | None = Query(None),
    q: str | None = Query(None),
    visibility: str | None = Query(None),  # private | org | system
    sort: str = Query(default="created_at"),   # created_at | title | loai_van_ban | ngay_ban_hanh
    order: str = Query(default="desc"),        # asc | desc
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if visibility == "private":
        base = select(ReferenceDocument).where(
            ReferenceDocument.created_by == current_user.id,
            ReferenceDocument.visibility == "private",
        )
    elif visibility == "org":
        base = select(ReferenceDocument).where(ReferenceDocument.visibility == "org")
    elif visibility == "system":
        base = select(ReferenceDocument).where(ReferenceDocument.visibility == "system")
    else:
        # Default: own documents (all visibilities)
        base = select(ReferenceDocument).where(ReferenceDocument.created_by == current_user.id)

    if loai:
        base = base.where(ReferenceDocument.loai_van_ban == loai)
    if hieu_luc:
        base = base.where(ReferenceDocument.hieu_luc == hieu_luc)
    if q:
        pattern = f"%{q}%"
        base = base.where(
            or_(
                ReferenceDocument.title.ilike(pattern),
                ReferenceDocument.trich_yeu.ilike(pattern),
                ReferenceDocument.so_ki_hieu.ilike(pattern),
                ReferenceDocument.co_quan_ban_hanh.ilike(pattern),
            )
        )

    total_result = await db.execute(select(sql_func.count()).select_from(base.subquery()))
    total = total_result.scalar() or 0

    sort_col = _SORT_COLUMNS.get(sort, ReferenceDocument.created_at)
    sort_expr = sort_col.asc() if order == "asc" else sort_col.desc()
    docs = (await db.execute(
        base.order_by(sort_expr).offset(skip).limit(limit)
    )).scalars().all()

    # Batch-count chunks for all docs in one query
    chunk_map: dict[str, int] = {}
    if docs:
        counts = await db.execute(
            select(ReferenceDocChunk.document_id, sql_func.count(ReferenceDocChunk.id))
            .where(ReferenceDocChunk.document_id.in_([d.id for d in docs]))
            .group_by(ReferenceDocChunk.document_id)
        )
        chunk_map = {row[0]: row[1] for row in counts.all()}

    items = [_add_url(d, chunk_count=chunk_map.get(d.id, 0)) for d in docs]
    return RefDocListResponse(items=items, total=total, skip=skip, limit=limit)


# ── Semantic search — MUST be before /{doc_id} ───────────────────────────────

@router.get("/search", response_model=RefDocSearchResponse)
async def search_ref_docs(
    q: str = Query(..., min_length=1),
    limit: int = Query(5, ge=1, le=20),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from app.services import embedding_service

    if not embedding_service.is_available():
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Embedding model not loaded — semantic search unavailable",
        )

    # Embed query in thread (CPU-bound)
    query_vector = await asyncio.to_thread(embedding_service.embed_text, q)

    # pgvector cosine distance: <=> operator
    # score = 1 - cosine_distance  (higher = more similar)
    stmt = (
        select(
            ReferenceDocument,
            (1 - ReferenceDocument.embedding.cosine_distance(query_vector)).label("score"),
        )
        .where(ReferenceDocument.created_by == current_user.id)
        .where(ReferenceDocument.embedding.is_not(None))
        .order_by(ReferenceDocument.embedding.cosine_distance(query_vector))
        .limit(limit)
    )

    result = await db.execute(stmt)
    rows = result.all()

    items = [_add_url(doc, float(score)) for doc, score in rows]
    return RefDocSearchResponse(items=items, query=q)


# ── Chunk-level semantic search — MUST be before /{doc_id} ──────────────────

@router.get("/search/chunks", response_model=RefDocChunkSearchResponse)
async def search_ref_doc_chunks(
    q: str = Query(..., min_length=1),
    limit: int = Query(5, ge=1, le=20),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from app.services import embedding_service
    from app.models.reference_doc_chunk import ReferenceDocChunk

    if not embedding_service.is_available():
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Embedding model not loaded — semantic search unavailable",
        )

    query_vector = await asyncio.to_thread(embedding_service.embed_text, q)

    stmt = (
        select(
            ReferenceDocChunk,
            ReferenceDocument.title,
            ReferenceDocument.so_ki_hieu,
            (1 - ReferenceDocChunk.embedding.cosine_distance(query_vector)).label("score"),
        )
        .join(ReferenceDocument, ReferenceDocChunk.document_id == ReferenceDocument.id)
        .where(ReferenceDocument.created_by == current_user.id)
        .where(ReferenceDocChunk.embedding.is_not(None))
        .order_by(ReferenceDocChunk.embedding.cosine_distance(query_vector))
        .limit(limit)
    )

    result = await db.execute(stmt)
    rows = result.all()

    items = [
        RefDocChunkSearchItem(
            document_id=chunk.document_id,
            document_title=title,
            so_ki_hieu=so_ki_hieu,
            chunk_index=chunk.chunk_index,
            dieu_khoan=chunk.dieu_khoan,
            content_preview=chunk.content[:200],
            score=float(score),
        )
        for chunk, title, so_ki_hieu, score in rows
    ]
    return RefDocChunkSearchResponse(items=items, query=q)


# ── Full-text search (Vietnamese, unaccent) — MUST be before /{doc_id} ───────

@router.get("/search/fulltext", response_model=RefDocFTSResponse)
async def search_ref_docs_fulltext(
    q: str = Query(..., min_length=1),
    limit: int = Query(10, ge=1, le=50),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # Build tsquery: "ho tich" → "ho & tich"
    terms = [t.strip() for t in q.strip().split() if t.strip()]
    tsquery_str = " & ".join(terms)

    tsq_expr = sql_func.to_tsquery("simple", sql_func.unaccent(tsquery_str))
    rank_expr = sql_func.ts_rank(ReferenceDocument.search_vector, tsq_expr).label("rank")

    stmt = (
        select(ReferenceDocument, rank_expr)
        .where(ReferenceDocument.created_by == current_user.id)
        .where(ReferenceDocument.search_vector.op("@@")(tsq_expr))
        .order_by(rank_expr.desc())
        .limit(limit)
    )

    result = await db.execute(stmt)
    rows = result.all()

    items = []
    for doc, rank in rows:
        resp = RefDocFTSItem.model_validate(doc)
        if doc.file_path:
            try:
                resp.download_url = get_file_url(doc.file_path, bucket_name=REF_DOCS_BUCKET)
            except Exception:
                pass
        resp.rank = float(rank)
        items.append(resp)

    return RefDocFTSResponse(items=items, query=q)


# ── Metadata preview (polling after upload) ───────────────────────────────────

@router.get("/{doc_id}/metadata-preview", response_model=MetadataPreviewResponse)
async def get_metadata_preview_endpoint(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # Verify ownership
    result = await db.execute(
        select(ReferenceDocument).where(
            ReferenceDocument.id == doc_id,
            ReferenceDocument.created_by == current_user.id,
        )
    )
    if not result.scalar_one_or_none():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    from app.core.redis import get_redis
    from app.services.metadata_extraction_service import get_metadata_preview
    redis = await get_redis()
    meta = await get_metadata_preview(doc_id, redis)

    if meta is None:
        return MetadataPreviewResponse(doc_id=doc_id, status="not_available")

    confidence_data = meta.pop("confidence", {})
    confidence = MetadataConfidence(**confidence_data) if confidence_data else MetadataConfidence()

    return MetadataPreviewResponse(
        doc_id=doc_id,
        status="ready",
        fields=meta,
        confidence=confidence,
    )


# ── Metadata confirm (user edits + writes to DB) ──────────────────────────────

@router.post("/{doc_id}/metadata-confirm", response_model=RefDocResponse)
async def confirm_metadata(
    doc_id: str,
    body: MetadataConfirmRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(ReferenceDocument).where(
            ReferenceDocument.id == doc_id,
            ReferenceDocument.created_by == current_user.id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Update only non-None fields (can_cu has no DB column — skipped)
    updatable = ("so_ki_hieu", "ngay_ban_hanh", "co_quan_ban_hanh",
                 "nguoi_ky", "trich_yeu", "hieu_luc", "tom_tat")
    for field in updatable:
        value = getattr(body, field)
        if value is not None:
            setattr(doc, field, value)

    await db.flush()
    await db.refresh(doc)

    # Remove Redis preview cache now that data is confirmed in DB
    from app.core.redis import get_redis
    redis = await get_redis()
    await redis.delete(f"metadata_preview:{doc_id}")

    return _add_url(doc)


# ── Batch upload (AI auto-extract metadata) ──────────────────────────────────

@router.post("/upload-batch", status_code=status.HTTP_202_ACCEPTED)
async def upload_ref_batch(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    visibility: str = Form(default="private"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload nhiều file cùng lúc. AI tự trích xuất metadata. Trả 202 ngay."""
    settings = get_settings()
    if len(files) > settings.upload_max_files:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Tối đa {settings.upload_max_files} file mỗi lần upload",
        )
    if visibility not in ("private", "org", "system"):
        visibility = "private"

    redis = await get_redis()
    jobs = []

    for file in files:
        file_data = await file.read()
        filename = file.filename or "unknown"
        file_type = file.content_type or "application/octet-stream"
        doc_id = str(uuid.uuid4())
        job_id = str(uuid.uuid4())

        # Upload to MinIO
        object_name = f"{current_user.id}/{doc_id}/{uuid.uuid4()}_{filename}"

        def _upload(data=file_data, name=object_name, ct=file_type):
            client = get_minio_client()
            ensure_bucket_exists(client, REF_DOCS_BUCKET)
            upload_file_data(client, REF_DOCS_BUCKET, name, data, ct)

        await asyncio.get_event_loop().run_in_executor(None, _upload)

        # Create ReferenceDocument with placeholder values (LLM fills later)
        doc = ReferenceDocument(
            id=doc_id,
            title=filename,
            loai_van_ban="",
            so_ki_hieu="",
            co_quan_ban_hanh="",
            trich_yeu="",
            hieu_luc="chua",
            visibility=visibility,
            file_path=object_name,
            file_size=len(file_data),
            file_type=file_type,
            created_by=current_user.id,
        )
        db.add(doc)

        # Redis job status
        await redis.set(
            f"ref_job:{job_id}",
            json.dumps({"status": "pending", "filename": filename, "doc_id": doc_id}),
            ex=settings.doc_job_ttl_seconds,
        )

        background_tasks.add_task(
            process_reference_background,
            job_id=job_id,
            doc_id=doc_id,
            file_data=file_data,
            filename=filename,
            file_type=file_type,
            owner_id=str(current_user.id),
        )

        jobs.append({"job_id": job_id, "filename": filename})
        logger.info("[ref_batch] queued job=%s doc=%s file=%s", job_id, doc_id, filename)

    await db.flush()
    return {"jobs": jobs}


@router.get("/status/{job_id}")
async def get_ref_job_status(
    job_id: str,
    current_user: User = Depends(get_current_user),
):
    """Poll trạng thái xử lý của một batch upload job (ref doc)."""
    redis = await get_redis()
    data = await redis.get(f"ref_job:{job_id}")
    if not data:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Job not found or expired")
    info = json.loads(data)
    return {
        "job_id": job_id,
        "status": info.get("status", "unknown"),
        "filename": info.get("filename", ""),
        "doc_id": info.get("doc_id", None),
        "error": info.get("error", None),
    }


# ── Create ───────────────────────────────────────────────────────────────────

@router.post("/", response_model=RefDocResponse, status_code=status.HTTP_201_CREATED)
async def create_ref_doc(
    doc_in: RefDocCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = ReferenceDocument(**doc_in.model_dump(), created_by=current_user.id)
    db.add(doc)
    await db.flush()
    await db.refresh(doc)
    return _add_url(doc)


# ── Get one ──────────────────────────────────────────────────────────────────

@router.get("/{doc_id}", response_model=RefDocResponse)
async def get_ref_doc(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(ReferenceDocument).where(
            ReferenceDocument.id == doc_id,
            ReferenceDocument.created_by == current_user.id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    return _add_url(doc)


# ── Update ───────────────────────────────────────────────────────────────────

@router.put("/{doc_id}", response_model=RefDocResponse)
async def update_ref_doc(
    doc_id: str,
    doc_in: RefDocUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(ReferenceDocument).where(
            ReferenceDocument.id == doc_id,
            ReferenceDocument.created_by == current_user.id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    for field, value in doc_in.model_dump(exclude_unset=True).items():
        setattr(doc, field, value)

    await db.flush()
    await db.refresh(doc)
    return _add_url(doc)


# ── Update hieu_luc ──────────────────────────────────────────────────────────

@router.patch("/{doc_id}/hieu-luc", response_model=RefDocResponse)
async def update_hieu_luc(
    doc_id: str,
    body: HieuLucUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if body.hieu_luc not in _VALID_HIEU_LUC:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"hieu_luc phải là một trong: {sorted(_VALID_HIEU_LUC)}",
        )

    result = await db.execute(
        select(ReferenceDocument).where(
            ReferenceDocument.id == doc_id,
            ReferenceDocument.created_by == current_user.id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    old_val = doc.hieu_luc
    doc.hieu_luc = body.hieu_luc
    await db.flush()
    await db.refresh(doc)

    logger.info(
        "[hieu_luc] doc %s: %s → %s%s",
        doc_id, old_val, body.hieu_luc,
        f" ({body.ghi_chu})" if body.ghi_chu else "",
    )

    return _add_url(doc)


# ── Delete ───────────────────────────────────────────────────────────────────

@router.delete("/{doc_id}")
async def delete_ref_doc(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(ReferenceDocument).where(
            ReferenceDocument.id == doc_id,
            ReferenceDocument.created_by == current_user.id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    if doc.file_path:
        try:
            client = get_minio_client()
            delete_file_data(client, REF_DOCS_BUCKET, doc.file_path)
            logger.info("[ref_doc] deleted storage object: %s", doc.file_path)
        except Exception as e:
            logger.warning(
                "[ref_doc] MinIO delete failed for %s: %s — continuing with DB delete",
                doc.file_path, e,
            )

    await db.delete(doc)
    await db.commit()

    return {"status": "deleted", "doc_id": doc_id}


# ── Upload file + trigger embedding pipeline ─────────────────────────────────

@router.post("/{doc_id}/upload", response_model=RefDocResponse)
async def upload_ref_doc_file(
    doc_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(ReferenceDocument).where(
            ReferenceDocument.id == doc_id,
            ReferenceDocument.created_by == current_user.id,
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    file_data = await file.read()
    object_name = f"{current_user.id}/{doc_id}/{uuid.uuid4()}_{file.filename}"
    content_type = file.content_type or "application/octet-stream"

    # Upload to storage in a thread (blocking I/O)
    def _upload():
        client = get_minio_client()
        ensure_bucket_exists(client, REF_DOCS_BUCKET)
        upload_file_data(client, REF_DOCS_BUCKET, object_name, file_data, content_type)

    await asyncio.get_event_loop().run_in_executor(None, _upload)

    old_path = doc.file_path
    doc.file_path = object_name
    doc.file_size = len(file_data)
    doc.file_type = content_type

    try:
        await db.flush()
        await db.refresh(doc)
    except Exception:
        # DB update failed — clean up the just-uploaded MinIO object
        def _cleanup():
            try:
                client = get_minio_client()
                delete_file_data(client, REF_DOCS_BUCKET, object_name)
            except Exception:
                pass
        await asyncio.get_event_loop().run_in_executor(None, _cleanup)
        raise

    # Remove previous file only after DB flush succeeds
    if old_path and old_path != object_name:
        def _remove_old():
            try:
                client = get_minio_client()
                delete_file_data(client, REF_DOCS_BUCKET, old_path)
            except Exception:
                pass
        await asyncio.get_event_loop().run_in_executor(None, _remove_old)

    # Kick off embedding in the background — does not block the response
    from app.services.pipeline_service import process_document_embedding
    background_tasks.add_task(process_document_embedding, doc_id)

    return _add_url(doc)


# ── Content (chunks for OCR viewer) ──────────────────────────────────────────

@router.get("/{doc_id}/content", response_model=RefDocContentResponse)
async def get_reference_doc_content(
    doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Trả về toàn bộ chunks của một văn bản tham chiếu, ghép theo chunk_index."""
    doc_result = await db.execute(
        select(ReferenceDocument).where(
            ReferenceDocument.id == doc_id,
            ReferenceDocument.created_by == current_user.id,
        )
    )
    doc = doc_result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    chunks_result = await db.execute(
        select(ReferenceDocChunk)
        .where(ReferenceDocChunk.document_id == doc_id)
        .order_by(ReferenceDocChunk.chunk_index.asc())
    )
    chunks = chunks_result.scalars().all()

    return RefDocContentResponse(
        id=doc.id,
        title=doc.title,
        so_ki_hieu=doc.so_ki_hieu if doc.so_ki_hieu else None,
        loai_van_ban=doc.loai_van_ban if doc.loai_van_ban else None,
        created_at=doc.created_at,
        chunks=[
            ChunkItem(
                chunk_index=c.chunk_index,
                content=c.content,
                dieu_khoan=c.dieu_khoan,
            )
            for c in chunks
        ],
    )


# ── Export (DOCX or PDF) ──────────────────────────────────────────────────────

@router.get("/{doc_id}/export")
async def export_reference_doc(
    doc_id: str,
    format: str = Query("docx", pattern="^(docx|pdf)$"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Xuất văn bản tham chiếu ra file DOCX hoặc PDF."""
    doc_result = await db.execute(
        select(ReferenceDocument).where(
            ReferenceDocument.id == doc_id,
            ReferenceDocument.created_by == current_user.id,
        )
    )
    doc = doc_result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    chunks_result = await db.execute(
        select(ReferenceDocChunk)
        .where(ReferenceDocChunk.document_id == doc_id)
        .order_by(ReferenceDocChunk.chunk_index.asc())
    )
    chunks = chunks_result.scalars().all()
    if not chunks:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Văn bản chưa được lập chỉ mục",
        )

    full_text = "\n\n".join(c.content for c in chunks)
    safe_filename = re.sub(r'[/\\:*?"<>|]', "_", doc.so_ki_hieu or doc.title or doc_id)[:120]

    if format == "docx":
        from docx import Document as DocxDocument
        docx_doc = DocxDocument()
        docx_doc.add_heading(doc.title, level=1)
        if doc.so_ki_hieu:
            docx_doc.add_paragraph(f"Số ký hiệu: {doc.so_ki_hieu}")
        docx_doc.add_paragraph("")
        docx_doc.add_paragraph(full_text)
        buf = io.BytesIO()
        docx_doc.save(buf)
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_filename}.docx"'},
        )

    # format == "pdf"
    import html as _html_esc
    from app.services.pdf_service import _write_pdf, _ensure_fonts, _build_css
    font = _ensure_fonts()
    css = _build_css(font)
    title_h = _html_esc.escape(doc.title)
    text_h = _html_esc.escape(full_text)
    meta_html = (
        f'<p style="font-size:12pt;color:#555;margin-bottom:8pt;">'
        f'Số ký hiệu: {_html_esc.escape(doc.so_ki_hieu)}</p>\n'
        if doc.so_ki_hieu else ""
    )
    html_str = (
        "<!DOCTYPE html>\n<html lang='vi'>\n<head>\n"
        "  <meta charset='UTF-8'>\n"
        f"  <style>{css}\n"
        f"    h1 {{ font-size: 15pt; font-weight: bold; text-align: center; margin-bottom: 4pt; }}\n"
        f"    pre {{ white-space: pre-wrap; font-family: '{font}', serif; font-size: 13pt; line-height: 1.6; }}\n"
        "  </style>\n</head>\n<body>\n"
        f"  <h1>{title_h}</h1>\n"
        f"{meta_html}"
        f"  <pre>{text_h}</pre>\n"
        "</body>\n</html>"
    )
    pdf_bytes = await asyncio.to_thread(_write_pdf, html_str)
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_filename}.pdf"'},
    )
