from fastapi import APIRouter, BackgroundTasks, Body, Depends, HTTPException, Query, status, UploadFile, File
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func as sql_func, case
from datetime import timedelta
from typing import List, Optional
from app.core.config import get_settings
from app.core.database import get_db
from app.core.redis import get_redis
from app.core.storage import upload_file
from app.api.deps import get_current_user
from app.models.user import User
from app.models.document import Document
from app.models.trich_yeu_history import TrichYeuHistory
from app.schemas.document import DocumentCreate, DocumentResponse, DocumentUpdate
from app.services.document_pipeline_service import process_document_background
from datetime import datetime
import asyncio
import json
import logging
import re
import uuid
from urllib.parse import quote

logger = logging.getLogger(__name__)

router = APIRouter()

# Prefix theo NĐ30 cho từng loại văn bản (abbreviation key)
_TRICH_YEU_PREFIX: dict[str, str] = {
    "CV": "V/v ",
    "TB": "Về việc ",
    "TTr": "Về việc ",
    "BC": "Về ",
    "GM": "Về việc ",
    "GGT": "Về việc ",
}


def _strip_html(html: str) -> str:
    return re.sub(r"<[^>]+>", " ", html).strip()


def _extract_tiptap_text(content_str: str) -> str:
    """Extract plain text from document.content — handles nd30 JSON and TipTap JSON."""
    if not content_str:
        return ""
    try:
        data = json.loads(content_str)
    except Exception:
        return _strip_html(content_str)

    # TipTap native format: {"type": "doc", "content": [...]}
    if data.get("type") == "doc":
        texts: list[str] = []
        def _walk(node: dict) -> None:
            if node.get("type") == "text":
                texts.append(node.get("text", ""))
            for child in (node.get("content") or []):
                _walk(child)
        _walk(data)
        return " ".join(t for t in texts if t).strip()

    # nd30 / nd30-compatible format
    parts: list[str] = []
    if data.get("trichYeu"):
        parts.append(f"Trích yếu: {data['trichYeu']}")
    if data.get("canCu"):
        parts.append(f"Căn cứ:\n{_strip_html(data['canCu'])}")
    if data.get("noiDung"):
        parts.append(f"Nội dung:\n{_strip_html(data['noiDung'])}")
    noi_nhan = data.get("noiNhan")
    if noi_nhan:
        if isinstance(noi_nhan, list):
            parts.append(f"Nơi nhận:\n{chr(10).join(noi_nhan)}")
        elif isinstance(noi_nhan, str):
            parts.append(f"Nơi nhận:\n{noi_nhan}")
    if parts:
        return "\n\n".join(parts)

    # Fallback: strip HTML from raw string
    return _strip_html(content_str)


def _extract_file_text(data: bytes, file_type: str | None, filename: str) -> str:
    """Synchronous text extraction from DOCX or PDF bytes. Call via asyncio.to_thread."""
    import io as _io
    ft = (file_type or "").lower()
    fn = (filename or "").lower()
    try:
        if "pdf" in ft or fn.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(_io.BytesIO(data)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
            return "\n\n".join(pages).strip()
        if "word" in ft or "docx" in ft or fn.endswith(".docx"):
            from docx import Document as _DocxDoc
            doc = _DocxDoc(_io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:
        logger.warning("[upload] text extraction failed for %s: %s", filename, exc)
    return ""


def _normalize_trich_yeu(trich_yeu: str, loai_vb: str | None) -> str:
    t = trich_yeu.strip()
    if not t:
        return t
    if t.startswith(("V/v", "Về việc", "Về ")):
        return t
    prefix = _TRICH_YEU_PREFIX.get(loai_vb or "", "")
    return f"{prefix}{t.lstrip()}" if prefix else t


async def _save_trich_yeu_history(
    doc: Document,
    db: AsyncSession,
    user_id: str,
) -> None:
    try:
        content = json.loads(doc.content or "{}")
        trich_yeu = content.get("trichYeu", "").strip()
        loai_vb = content.get("loaiVanBan", doc.loai_vb or "VB")
    except (json.JSONDecodeError, AttributeError):
        return

    if not trich_yeu or len(trich_yeu) < 3:
        return

    normalized = _normalize_trich_yeu(trich_yeu, loai_vb)

    existing = (await db.execute(
        select(TrichYeuHistory).where(
            TrichYeuHistory.loai_van_ban == loai_vb,
            TrichYeuHistory.trich_yeu == normalized,
            TrichYeuHistory.created_by == user_id,
        )
    )).scalar_one_or_none()

    if existing:
        existing.used_count += 1
        existing.last_used_at = datetime.utcnow()
    else:
        db.add(TrichYeuHistory(
            loai_van_ban=loai_vb,
            trich_yeu=normalized,
            created_by=user_id,
            source_doc_id=str(doc.id),
        ))
    logger.debug("[trich_yeu_history] saved: %s / %s", loai_vb, normalized[:50])


@router.get("/")
async def list_documents(
    skip: int = 0,
    limit: int = 20,
    source: str | None = None,
    q: str | None = Query(None),
    loai_vb: str | None = Query(None),
    sort_by: str = Query("created_at"),
    sort_order: str = Query("desc"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    base = select(Document).where(Document.owner_id == current_user.id)
    if source in ("editor", "upload"):
        base = base.where(Document.source == source)
    if q:
        base = base.where(Document.title.ilike(f"%{q}%"))
    if loai_vb:
        base = base.where(Document.loai_vb == loai_vb)

    total_result = await db.execute(
        select(sql_func.count()).select_from(base.subquery())
    )
    total = total_result.scalar() or 0

    _sort_cols = {
        "title":      Document.title,
        "loai_vb":    Document.loai_vb,
        "created_at": Document.created_at,
    }
    sort_col = _sort_cols.get(sort_by, Document.created_at)
    sort_expr = sort_col.asc() if sort_order == "asc" else sort_col.desc()

    result = await db.execute(
        base.order_by(sort_expr).offset(skip).limit(limit)
    )
    items = result.scalars().all()
    return {"items": items, "total": total}


@router.post("/", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def create_document(
    document_in: DocumentCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = Document(
        **document_in.model_dump(),
        owner_id=current_user.id,
        source="editor",
    )
    db.add(document)
    await db.flush()
    await db.refresh(document)
    return document


# Must be before /{document_id} to avoid route conflict
@router.get("/stats")
async def get_document_stats(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Thống kê tài liệu của người dùng."""
    uid = current_user.id

    # Aggregate counts in one query
    agg = (await db.execute(
        select(
            sql_func.count(Document.id).label("total"),
            sql_func.count(case((Document.source == "editor", 1))).label("editor_count"),
            sql_func.count(case((Document.source == "upload", 1))).label("upload_count"),
        ).where(Document.owner_id == uid)
    )).one()

    # By loai_vb
    type_rows = (await db.execute(
        select(Document.loai_vb, sql_func.count(Document.id))
        .where(Document.owner_id == uid, Document.loai_vb.isnot(None))
        .group_by(Document.loai_vb)
    )).all()
    by_type = {r[0]: r[1] for r in type_rows}

    # Recent 7 days
    cutoff = datetime.utcnow() - timedelta(days=7)
    recent_7_days = (await db.execute(
        select(sql_func.count(Document.id))
        .where(Document.owner_id == uid, Document.created_at >= cutoff)
    )).scalar() or 0

    return {
        "total": agg.total,
        "editor_count": agg.editor_count,
        "upload_count": agg.upload_count,
        "by_type": by_type,
        "recent_7_days": recent_7_days,
    }


_LOAI_TO_ABBR: dict[str, str] = {
    "Quyết định": "QĐ", "Công văn": "CV", "Báo cáo": "BC",
    "Hướng dẫn": "HD", "Tờ trình": "TTr", "Thông báo": "TB",
    "Nghị quyết": "NQ", "Kế hoạch": "KH", "Chỉ thị": "CT",
}

_GENERATE_SYSTEM = """\
Bạn là chuyên gia soạn thảo văn bản hành chính Việt Nam theo Nghị định 30/2020/NĐ-CP.

== THÔNG TIN BẮT BUỘC ==
Loại văn bản: {loai_van_ban} (ký hiệu: {abbr}) — KHÔNG được tạo loại văn bản khác.
Năm hiện tại: {current_year} — số ký hiệu PHẢI dùng năm {current_year}.
Yêu cầu: {yeu_cau}

== TÀI LIỆU THAM CHIẾU ==
{context}

== QUY TẮC NGHIÊM NGẶT ==
1. LOẠI VĂN BẢN: field "loaiVanBan" PHẢI là "{abbr}". QUAN TRỌNG: không được thay bằng loại khác.
2. SỐ KÝ HIỆU: format [số]/{current_year}/{abbr}-[viết tắt cơ quan]. Ví dụ: 05/{current_year}/{abbr}-SKHCN
3. CĂN CỨ: CHỈ trích dẫn văn bản có trong TÀI LIỆU THAM CHIẾU. Nếu không có → ghi "[Cần bổ sung căn cứ pháp lý]". TUYỆT ĐỐI không tự bịa số hiệu văn bản.
4. NƠI NHẬN: Phải có ít nhất 2-3 đơn vị cụ thể phù hợp nội dung. Không chỉ "Như trên" và "Lưu VT".
5. CHỮ KÝ: Điền chức danh đúng: Sở/Cục → "GIÁM ĐỐC", UBND → "CHỦ TỊCH", Bộ → "BỘ TRƯỞNG", cơ quan khác → chức danh phù hợp.

== THỂ THỨC TRÌNH BÀY (NĐ30/2020/NĐ-CP) ==
- Font chữ: Times New Roman, cỡ 13-14pt, màu đen.
- Lề trang: trên 20-25mm, dưới 20-25mm, trái 30-35mm, phải 15-20mm.
- Dùng thông tin này để căn chỉnh độ dài nội dung phù hợp với khổ A4.

== CẤU TRÚC NỘI DUNG THEO LOẠI VĂN BẢN ==
Báo cáo (BC):
  <p><strong>I. TÌNH HÌNH THỰC HIỆN</strong></p><p>...</p><p><strong>II. KẾT QUẢ ĐẠT ĐƯỢC</strong></p><p>...</p><p><strong>III. KHÓ KHĂN, VƯỚNG MẮC</strong></p><p>...</p><p><strong>IV. KIẾN NGHỊ, ĐỀ XUẤT</strong></p><p>...</p>
Công văn (CV):
  <p>Kính gửi: [Cơ quan]</p><p>[Nội dung chính]</p><p>Kính đề nghị [cơ quan] [hành động].</p>
Tờ trình (TTr):
  <p>Kính trình [cấp trên]...</p><p><strong>I. SỰ CẦN THIẾT</strong></p><p>...</p><p><strong>II. NỘI DUNG ĐỀ XUẤT</strong></p><p>...</p><p><strong>III. KIẾN NGHỊ</strong></p><p>...</p>
Quyết định (QĐ):
  <p><strong>Điều 1:</strong> ...</p><p><strong>Điều 2:</strong> ...</p><p><strong>Điều 3:</strong> Quyết định này có hiệu lực kể từ ngày ký.</p>
Kế hoạch (KH):
  <p><strong>I. MỤC TIÊU</strong></p><p>...</p><p><strong>II. NỘI DUNG</strong></p><p>...</p><p><strong>III. TỔ CHỨC THỰC HIỆN</strong></p><p>...</p>
Thông báo (TB) / Hướng dẫn (HD): Nội dung theo đề mục phù hợp với chủ đề.

== OUTPUT ==
Trả về JSON hợp lệ (KHÔNG markdown, KHÔNG ```json, KHÔNG giải thích):
{{
  "loaiVanBan": "{abbr}",
  "soKyHieu": "05/{current_year}/{abbr}-[viết tắt cơ quan]",
  "trichYeu": "Trích yếu ngắn gọn 1 câu về nội dung",
  "canCu": "<p>Căn cứ [văn bản từ tài liệu tham chiếu hoặc để trống nếu không có]</p>",
  "noiDung": "<p>[Nội dung đúng cấu trúc loại {abbr} như hướng dẫn]</p>",
  "coQuanBanHanh": "[Tên cơ quan ban hành đầy đủ]",
  "chucVuKy": "[GIÁM ĐỐC / CHỦ TỊCH / BỘ TRƯỞNG tùy cơ quan]",
  "noiNhan": "- [Cơ quan nhận 1];\\n- [Cơ quan nhận 2];\\n- Lưu: VT, [đơn vị liên quan]."
}}\
"""


class GenerateDocumentRequest(BaseModel):
    document_id: str
    loai_van_ban: str       # "Quyết định" | "QĐ" | ...
    yeu_cau: str
    source_ids: list[str] = []


@router.post("/generate")
async def generate_document(
    body: GenerateDocumentRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Gọi LLM để soạn thảo văn bản từ mô tả + tài liệu tham chiếu."""
    from app.services.llm_service import llm_service
    from app.services.rag_service import rag_service

    # Resolve abbreviation
    abbr = _LOAI_TO_ABBR.get(body.loai_van_ban, body.loai_van_ban)
    display = body.loai_van_ban if body.loai_van_ban not in _LOAI_TO_ABBR.values() \
        else next((k for k, v in _LOAI_TO_ABBR.items() if v == body.loai_van_ban), body.loai_van_ban)

    # Verify document ownership
    doc_result = await db.execute(
        select(Document).where(
            Document.id == body.document_id,
            Document.owner_id == current_user.id,
        )
    )
    document = doc_result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # If LLM not available → skip
    if not llm_service._base_url:
        logger.warning("[generate] LLM not configured — skipping")
        return {"status": "skipped", "document_id": body.document_id}

    # Build RAG context from source_ids (if any)
    context = ""
    if body.source_ids:
        try:
            chunks = await rag_service.retrieve(
                body.yeu_cau, db, top_k=5, min_score=0.2,
                source_ids=body.source_ids,
            )
            context = rag_service.build_context(chunks) if chunks else ""
        except Exception as exc:
            logger.warning("[generate] RAG retrieval failed: %s", exc)

    # Build prompt and call LLM
    system_prompt = _GENERATE_SYSTEM.format(
        loai_van_ban=display,
        yeu_cau=body.yeu_cau,
        context=context or "Không có tài liệu tham chiếu.",
        abbr=abbr,
        current_year=datetime.now().year,
    )

    try:
        raw = await llm_service.chat(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": body.yeu_cau},
            ],
            temperature=0.2,
            max_tokens=4000,
            json_mode=True,
        )
        raw_clean = sanitize_json_string(raw)
        generated = json.loads(raw_clean)
    except Exception as exc:
        logger.error("[generate] LLM call failed: %s", exc)
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                            detail=f"LLM error: {exc}")

    # Ensure loaiVanBan is abbreviation
    generated["loaiVanBan"] = abbr
    generated["ai_generated"] = True

    # Normalize noiNhan: LLM sometimes returns string instead of array
    noi_nhan = generated.get("noiNhan", [])
    if isinstance(noi_nhan, str):
        lines = [ln.strip() for ln in noi_nhan.replace("\\n", "\n").split("\n") if ln.strip()]
        generated["noiNhan"] = lines if lines else []
    elif not isinstance(noi_nhan, list):
        generated["noiNhan"] = []

    # Save to document
    content_str = json.dumps({"version": "nd30", **generated}, ensure_ascii=False)
    document.content = content_str
    document.title = generated.get("trichYeu") or body.yeu_cau[:100] or "Văn bản mới"
    document.loai_vb = abbr
    await db.flush()

    logger.info("[generate] doc=%s loai=%s", body.document_id, abbr)
    return {"status": "done", "document_id": body.document_id, "content": generated}


@router.get("/next-number")
async def get_next_number(
    loai: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    nam = datetime.now().year
    result = await db.execute(
        select(sql_func.count(Document.id)).where(
            Document.loai_vb == loai,
            Document.nam == nam,
            Document.owner_id == current_user.id,
        )
    )
    count = result.scalar() or 0
    return {"so": count + 1, "nam": nam, "loai": loai}


@router.post("/upload-batch", status_code=status.HTTP_202_ACCEPTED)
async def upload_batch(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload nhiều file cùng lúc. Trả về 202 ngay, xử lý nền theo từng job."""
    settings = get_settings()
    if len(files) > settings.upload_max_files:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Tối đa {settings.upload_max_files} file mỗi lần upload",
        )

    redis = await get_redis()
    jobs = []

    for file in files:
        file_data = await file.read()
        job_id = str(uuid.uuid4())
        doc_id = str(uuid.uuid4())
        filename = file.filename or "unknown"
        file_type = file.content_type or "application/octet-stream"

        # Upload file to MinIO immediately
        object_name = f"{current_user.id}/{doc_id}/{uuid.uuid4()}_{filename}"
        await upload_file(file_data, object_name, file_type)

        # Create Document record
        document = Document(
            id=doc_id,
            title=filename,
            owner_id=current_user.id,
            file_path=object_name,
            file_type=file_type,
            source="upload",
        )
        db.add(document)

        # Save initial job status in Redis (TTL 24h)
        await redis.set(
            f"doc_job:{job_id}",
            json.dumps({"status": "pending", "filename": filename, "doc_id": doc_id}),
            ex=settings.doc_job_ttl_seconds,
        )

        # Schedule background processing (file bytes already read — safe after response)
        background_tasks.add_task(
            process_document_background,
            job_id=job_id,
            doc_id=doc_id,
            file_data=file_data,
            filename=filename,
            file_type=file_type,
            owner_id=str(current_user.id),
        )

        jobs.append({"job_id": job_id, "filename": filename})
        logger.info("[upload_batch] queued job=%s doc=%s file=%s", job_id, doc_id, filename)

    await db.flush()
    return {"jobs": jobs}


@router.get("/status/{job_id}")
async def get_job_status(
    job_id: str,
    current_user: User = Depends(get_current_user),
):
    """Poll trạng thái xử lý của một file upload job."""
    redis = await get_redis()
    data = await redis.get(f"doc_job:{job_id}")
    if not data:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Job not found or expired")
    info = json.loads(data)
    return {
        "job_id": job_id,
        "status": info.get("status", "unknown"),
        "filename": info.get("filename", ""),
        "error": info.get("error", None),
    }


@router.get("/test-docx")
async def test_docx_export(current_user: User = Depends(get_current_user)):
    """Test endpoint: generate DOCX với data mẫu cứng để isolate lỗi."""
    sample = {
        "loaiVanBan": "QĐ",
        "coQuanChuQuan": "UỶ BAN NHÂN DÂN TỈNH",
        "coQuanBanHanh": "SỞ GIÁO DỤC VÀ ĐÀO TẠO",
        "soKyHieu": "123/QĐ-SGD",
        "diaDanh": "Hà Nội",
        "ngayThang": "ngày 15 tháng 05 năm 2025",
        "trichYeu": "Quyết định về việc phê duyệt kế hoạch đào tạo năm 2025",
        "canCu": "<p>Căn cứ Luật Giáo dục số 43/2019/QH14 ngày 14/6/2019;</p>",
        "noiDung": "<p>Điều 1. Phê duyệt kế hoạch đào tạo năm học 2025-2026.</p><p>Điều 2. Quyết định này có hiệu lực kể từ ngày ký.</p>",
        "quyenHanKy": "KT. GIÁM ĐỐC",
        "chucDanhTapThe": "",
        "chucVuKy": "PHÓ GIÁM ĐỐC",
        "hoTenKy": "Nguyễn Văn An",
        "noiNhan": ["- Ban Giám đốc Sở", "- Các phòng ban trực thuộc", "- Lưu: VT, VP."],
    }
    try:
        from app.services.docx_service import generate_docx
        docx_bytes = await generate_docx(sample)
    except Exception as exc:
        logger.error("test-docx failed: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Test DOCX thất bại: {exc}")

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="test_vanban.docx"'},
    )


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(
            Document.id == document_id, Document.owner_id == current_user.id
        )
    )
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    return document


@router.patch("/{document_id}", response_model=DocumentResponse)
async def update_document(
    document_id: str,
    document_in: DocumentUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(
            Document.id == document_id, Document.owner_id == current_user.id
        )
    )
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    for field, value in document_in.model_dump(exclude_unset=True).items():
        setattr(document, field, value)

    await db.flush()
    await db.refresh(document)
    await _save_trich_yeu_history(document, db, current_user.id)
    return document


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(
            Document.id == document_id, Document.owner_id == current_user.id
        )
    )
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    await db.delete(document)


@router.post("/{document_id}/export/pdf")
async def export_document_pdf(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(
            Document.id == document_id, Document.owner_id == current_user.id
        )
    )
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    try:
        data = json.loads(document.content or "{}")
    except Exception:
        data = {}

    from app.services.pdf_service import generate_pdf
    pdf_bytes = await generate_pdf(data)

    # Build safe filename: {soKyHieu}_{title}.pdf
    so_ky = data.get("soKyHieu") or document.title or document_id
    raw_name = f"{so_ky}_{document.title or ''}".strip("_")
    safe_name = re.sub(r'[/\\:*?"<>|]', "-", raw_name)[:120] + ".pdf"
    encoded_name = quote(safe_name, safe="")

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=\"document.pdf\"; filename*=UTF-8''{encoded_name}"},
    )


@router.post("/{document_id}/export/docx")
async def export_document_docx(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(
            Document.id == document_id, Document.owner_id == current_user.id
        )
    )
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    try:
        data = json.loads(document.content or "{}")
    except Exception:
        data = {}

    try:
        from app.services.docx_service import generate_docx
        docx_bytes = await generate_docx(data)
    except Exception as exc:
        logger.error("export/docx failed for %s: %s", document_id, exc, exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Tạo DOCX thất bại: {exc}",
        )

    so_ky = data.get("soKyHieu") or document.title or document_id
    raw_name = f"{so_ky}_{document.title or ''}".strip("_")
    safe_name = re.sub(r'[/\\:*?"<>|]', "-", raw_name)[:120] + ".docx"
    encoded_name = quote(safe_name, safe="")

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"document.docx\"; filename*=UTF-8''{encoded_name}"},
    )


@router.post("/{document_id}/upload")
async def upload_document_file(
    document_id: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload file vào document editor: extract text đồng bộ, lưu nd30 JSON, trả extracted_text."""
    result = await db.execute(
        select(Document).where(
            Document.id == document_id, Document.owner_id == current_user.id
        )
    )
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    file_data = await file.read()
    filename = file.filename or "unknown"
    file_type = file.content_type or "application/octet-stream"
    object_name = f"{current_user.id}/{document_id}/{uuid.uuid4()}_{filename}"
    await upload_file(file_data, object_name, file_type)

    # Extract text synchronously in thread pool (avoids blocking event loop)
    extracted_text = await asyncio.to_thread(_extract_file_text, file_data, file_type, filename)
    logger.info("[upload] extracted %d chars from %s doc=%s", len(extracted_text), filename, document_id)

    # Save nd30-compatible content so AI Review and editor work immediately
    if extracted_text:
        paras = "".join(
            f"<p>{line}</p>"
            for line in extracted_text.splitlines()
            if line.strip()
        )
        document.content = json.dumps(
            {"version": "nd30", "noiDung": paras},
            ensure_ascii=False,
        )

    document.file_path = object_name
    document.file_type = file_type
    await db.flush()
    await db.refresh(document)

    return {
        "id": str(document.id),
        "title": document.title,
        "content": document.content,
        "file_path": document.file_path,
        "file_type": document.file_type,
        "loai_vb": document.loai_vb,
        "so_van_ban": document.so_van_ban,
        "nam": document.nam,
        "source": document.source,
        "owner_id": str(document.owner_id),
        "created_at": document.created_at.isoformat() if document.created_at else None,
        "updated_at": document.updated_at.isoformat() if document.updated_at else None,
        "extracted_text": extracted_text,
    }


_REVIEW_SYSTEM_PROMPT = """Bạn là chuyên gia rà soát văn bản hành chính Việt Nam \
theo Nghị định 30/2020/NĐ-CP về công tác văn thư.

NHIỆM VỤ: Phân tích văn bản và đề xuất chỉnh sửa theo 5 tiêu chí:
1. Chính tả và lỗi đánh máy (typo)
2. Thể thức văn bản theo NĐ30/2020 (quốc hiệu, tiêu ngữ, số ký hiệu, địa danh, ngày tháng, nơi nhận)
3. Văn phong hành chính (trang trọng, không dùng từ thông tục, câu rõ ràng)
4. Dấu câu và cách trình bày
5. Thuật ngữ pháp lý đúng chuẩn (viết tắt, tên cơ quan, chức danh)

NGUYÊN TẮC QUAN TRỌNG:
- CHỈ đề xuất sửa khi THỰC SỰ sai — không sửa nếu đã đúng
- KHÔNG thay đổi nội dung pháp lý, chỉ sửa hình thức và ngôn ngữ
- original và revised PHẢI KHÁC NHAU — bỏ qua nếu giống hệt
- Mỗi reason ngắn gọn, tối đa 15 từ
- Nếu văn bản đã chuẩn → trả về "changes": [] và summary tích cực

ĐỊNH DẠNG TRẢ VỀ (JSON thuần, KHÔNG markdown, KHÔNG dùng dấu ```):
{
  "reviewed_text": "toàn bộ văn bản đã sửa tất cả lỗi",
  "changes": [
    {
      "section": "trichYeu|canCu|noiDung|general",
      "type": "chinh_ta|the_thuc|van_phong|dau_cau|thuat_ngu",
      "original": "đoạn văn gốc (plain text, không HTML)",
      "revised": "đoạn văn đã sửa (plain text, không HTML)",
      "reason": "lý do ngắn gọn"
    }
  ],
  "summary": "tóm tắt 1-2 câu về chất lượng văn bản"
}"""


def sanitize_json_string(s: str) -> str:
    # Xóa control characters trừ \n \r \t
    s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)

    # Extract JSON nếu bị wrap trong markdown
    match = re.search(r'\{.*\}', s, re.DOTALL)
    if match:
        s = match.group(0)

    # Fix: escape newlines thật bên trong JSON string values
    def fix_newlines_in_strings(json_str: str) -> str:
        result = []
        in_string = False
        escape = False
        for char in json_str:
            if escape:
                result.append(char)
                escape = False
            elif char == '\\' and in_string:
                result.append(char)
                escape = True
            elif char == '"':
                result.append(char)
                in_string = not in_string
            elif char == '\n' and in_string:
                result.append('\\n')
            elif char == '\r' and in_string:
                result.append('\\r')
            elif char == '\t' and in_string:
                result.append('\\t')
            else:
                result.append(char)
        return ''.join(result)

    return fix_newlines_in_strings(s)


class ReviewRequest(BaseModel):
    content: Optional[str] = None


@router.post("/{document_id}/review")
async def review_document(
    document_id: str,
    body: ReviewRequest = Body(default=ReviewRequest()),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Gọi LLM rà soát chính tả/thể thức văn bản hành chính. Kết quả cache Redis 24h."""
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.owner_id == current_user.id,
        )
    )
    document = result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Ưu tiên content từ frontend (real-time), fallback đọc từ DB
    if body.content:
        plain_text = _extract_tiptap_text(body.content)
    else:
        plain_text = _extract_tiptap_text(document.content or "")

    if not plain_text.strip():
        raise HTTPException(
            status_code=422,
            detail="Vui lòng nhập nội dung văn bản trước khi rà soát",
        )

    try:
        from app.services.llm_service import llm_service
        if not llm_service._base_url:
            raise HTTPException(status_code=503, detail="LLM service không khả dụng")

        raw = await llm_service.chat(
            messages=[
                {"role": "system", "content": _REVIEW_SYSTEM_PROMPT},
                {"role": "user", "content": plain_text[:6000]},
            ],
            temperature=0.1,
            max_tokens=4000,
            json_mode=True,
        )
        logger.error("[review_doc] raw response: %r", raw[:200])
        raw_clean = sanitize_json_string(raw)
        review_data = json.loads(raw_clean)
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("[review_doc] doc=%s error: %s", document_id, exc)
        raise HTTPException(status_code=503, detail=f"LLM error: {exc}")

    redis = await get_redis()
    await redis.set(
        f"review:doc:{document_id}",
        json.dumps(review_data, ensure_ascii=False),
        ex=86400,
    )

    return {
        "doc_id": document_id,
        "reviewed_text": review_data.get("reviewed_text", ""),
        "changes": review_data.get("changes", []),
        "summary": review_data.get("summary", ""),
    }
